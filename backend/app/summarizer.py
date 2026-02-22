import os
import logging
from typing import List, Dict, Any, Optional
from langchain_core.documents import Document
from langchain.chains.summarize import load_summarize_chain
from langchain.prompts import PromptTemplate
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from app.config import llm

logger = logging.getLogger(__name__)

class LegalSummarizer:
    """Professional summarizer optimized for long Pakistani legal documents"""
    
    def __init__(self):
        self.llm = llm
        
    def summarize_long_document(self, docs: List[Document], summary_type: str = "executive") -> str:
        """Summarize long legal documents using Map-Reduce chain"""
        if not self.llm:
            return "Summarization service unavailable."
            
        if summary_type == "executive":
            combine_template = """
            Write a professional Executive Summary of the following legal decision/document.
            Focus on:
            1. The core legal issue.
            2. The evidence or precedents cited.
            3. The final ruling or decision.
            
            Keep it clear and easy for a layman to understand while maintaining legal accuracy.
            
            TEXT: {text}
            
            EXECUTIVE SUMMARY (Layman-friendly):"""
        else:
            combine_template = """
            Provide a Detailed Legal Analysis of the following document.
            Include:
            - Parties involved
            - Procedural history
            - Key arguments
            - Legal reasoning (Ratio Decidendi)
            - Final Order
            
            TEXT: {text}
            
            DETAILED ANALYSIS:"""
            
        map_template = "Summarize the following section of a legal document, preserving all key legal citations and parties mentioned: {text}"
        
        map_prompt = PromptTemplate(template=map_template, input_variables=["text"])
        combine_prompt = PromptTemplate(template=combine_template, input_variables=["text"])
        
        try:
            # Use Map-Reduce for long documents
            chain = load_summarize_chain(
                self.llm, 
                chain_type="map_reduce",
                map_prompt=map_prompt,
                combine_prompt=combine_prompt,
                verbose=False
            )
            
            summary = chain.run(docs)
            return summary
        except Exception as e:
            logger.error(f"Summarization failed: {e}")
            return f"Failed to generate summary: {str(e)}"

    def export_to_docx(self, content: str, title: str, output_path: str):
        """Export summary to a professional Word document"""
        doc = DocxDocument()
        doc.add_heading(title, 0)
        
        doc.add_paragraph("Legal AI - Professional Case Summary")
        doc.add_paragraph(f"Document: {title}")
        doc.add_paragraph("-" * 20)
        
        doc.add_paragraph(content)
        
        doc.save(output_path)
        return output_path

    def export_to_pdf(self, content: str, title: str, output_path: str):
        """Export summary to a professional PDF document"""
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Heading
        story.append(Paragraph(f"<b>{title}</b>", styles['Heading1']))
        story.append(Paragraph("<i>Legal AI Professional Case Summary</i>", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Content
        # Split by newlines and add paragraphs
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                story.append(Paragraph(line, styles['Normal']))
                story.append(Spacer(1, 6))
        
        doc.build(story)
        return output_path

summarizer = LegalSummarizer()
